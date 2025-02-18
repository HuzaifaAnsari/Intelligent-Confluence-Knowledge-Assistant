import pandas as pd
import uuid
import docx
import json
import re
import torch
from haystack.utils import ComponentDevice
from haystack.components.preprocessors import DocumentSplitter
from haystack.document_stores.types import DuplicatePolicy
from haystack_integrations.document_stores.elasticsearch import ElasticsearchDocumentStore
from haystack import Document
from haystack.components.preprocessors import DocumentCleaner
from haystack.components.embedders import SentenceTransformersDocumentEmbedder

from confluence.utils import confluence_program
from dotenv import load_dotenv
import os
load_dotenv()

elasticsearch_url =os.getenv('ELASTICSEARCH_URL')
elasticsearch_username=os.getenv('ELASTICSEARCH_USERNAME')
elasticsearch_password=os.getenv("ELASTICSEARCH_PASSWORD")
elasticsearch_indexname=os.getenv("ELASTICSEARCH_INDEXNAME")



def read_docx(doc):
    content = []

    for para in doc.paragraphs:
        content.append(para.text)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content.append(table_data)

    return content

def read_files(content):
    processed_data = {
        "tables": [],
        "queries": [],
        "data_types": [],
        "json_data": [],
        "hierarchical_data": []
    }

    for item in content:
        if isinstance(item, list):
            
            processed_data["tables"].append(item)
        elif re.match(r"SELECT .* FROM .*", item, re.IGNORECASE | re.DOTALL):
            
            processed_data["queries"].append(item)
        elif re.match(r"^\{.*\}$", item, re.DOTALL):
            
            try:
                json_data = json.loads(item)
                processed_data["json_data"].append(json_data)
            except json.JSONDecodeError:
                pass
        elif re.match(r"├──|└──", item):
            
            processed_data["hierarchical_data"].append(item)
        else:
            
            processed_data["data_types"].append(item)

    return processed_data




document_store = ElasticsearchDocumentStore(hosts = elasticsearch_url,basic_auth=(elasticsearch_username, elasticsearch_password), index=elasticsearch_indexname, embedding_similarity_function = "cosine",verify_certs=False)

response = confluence_program.query_search('type = page')
if response:
    page_texts = confluence_program.get_page_text(response)
    page_dict = {}
    for page_id, content in page_texts.items():
        plain_text = confluence_program.extract_plain_text(content['text'])
        
        page_dict[page_id] = {
            "title": content['title'],
            "text": confluence_program.text_to_docx(plain_text, f"{content['title']}.docx")
        }
            
data_list = []
for page_id, content in page_dict.items():
    data_list.append({
        "Page_ID": page_id,
        "Page_Title": content['title'],
        "Page": content['text']
    })

# Create a pandas DataFrame
df_1 = pd.DataFrame(data_list)


page_id, email, name, accountId, title, page_url,date = confluence_program.retrieve_confluence_pages()
page_data = [page_id, email, name, accountId, page_url,date]
pages_list_data = [lst[1:] for lst in page_data]

data=[]
for i in range(len(page_data)):
    data.append(page_data[i])

transposed_data = list(map(list, zip(*data)))
columns = ["Page_ID", "Author_Email", "Author_Name", "Author_ID", "Page_URL", "Date"]

df_2 = pd.DataFrame(transposed_data, columns=columns)


df = pd.merge(df_1, df_2, on="Page_ID", how="inner")

df["UUID"] = [str(uuid.uuid4()) for _ in range(len(df))]

col = ["UUID", "Page_ID", "Author_Email", "Author_Name", "Author_ID", "Page_Title", "Page_URL", "Date"]

files_path=list(df["Page"])

metadata_df=df[col]
metadata_list = metadata_df.to_dict(orient="records")
file_metadata_pairs = []
for i in range(len(files_path)):
    content=read_docx(files_path[i])
    processed_data = read_files(content)
    metadata = metadata_list[i]
    processed_data_str = json.dumps(processed_data, indent=2)  
    pair = {"content": processed_data_str, "meta": metadata}
    file_metadata_pairs.append(pair)

print(f"Len of Metadata Pair {len(file_metadata_pairs)}")

print(f"Checking First Pair of Data {file_metadata_pairs[1].get('content')}")
    
cleaner = DocumentCleaner(
    unicode_normalization="NFKC",  
    ascii_only=False,              
    remove_empty_lines=True,       
    remove_extra_whitespaces=True, 
    remove_repeated_substrings=False   
)
splitter = DocumentSplitter(
    split_by="word", 
    split_length=500, 
    split_overlap=50, 
    split_threshold=200
)

docs = [Document(content=pair["content"], meta=pair["meta"]) for pair in file_metadata_pairs]

cleaned_docs = cleaner.run(docs)
split_docs = splitter.run(cleaned_docs.get("documents"))

if torch.cuda.is_available():
    print("Using GPU")
    document_embedder = SentenceTransformersDocumentEmbedder(
        model="BAAI/bge-m3" ,device = ComponentDevice.from_str("cuda:0")
    )
else:
    print("Using CPU")
    document_embedder = SentenceTransformersDocumentEmbedder(
        model="BAAI/bge-m3" )
document_embedder.warm_up()
documents_with_embeddings = document_embedder.run(split_docs.get("documents"))
document_store.write_documents(documents_with_embeddings.get("documents"), policy=DuplicatePolicy.SKIP)