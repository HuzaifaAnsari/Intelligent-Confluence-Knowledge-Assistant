from dotenv import load_dotenv
import os
import requests
from requests.auth import HTTPBasicAuth
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
load_dotenv()

confluence_api_token = os.getenv('CONFLUENCE_API_TOKEN')
confluence_username = os.getenv('CONFLUENCE_USERNAME')
confluence_url = os.getenv('CONFLUENCE_URL')
confluence_space_key = os.getenv('SPACE_KEY')
confluence_user_email = os.getenv('USER_EMAIL')


AUTH = HTTPBasicAuth(confluence_user_email, confluence_api_token)
HEADERS = {
    "Content-Type": "application/json",
    "Accept": "application/json",
}

def list_spaces():
    url = f"{confluence_url}/spaces"
    response = requests.get(url, headers=HEADERS, auth=AUTH)
    if response.status_code == 200:
        print(json.dumps(json.loads(response.text), sort_keys=True, indent=4, separators=(",", ": ")))
    else:
        print(f"Failed to retrieve spaces. Status code: {response.status_code}")
        print(response.text)

def delete_page(page_id: str) -> None:
    url = f"{confluence_url}/content/{page_id}"
    response = requests.delete(url, headers=HEADERS, auth=AUTH)
    if response.status_code == 204:
        print(f"Page with ID {page_id} deleted successfully!")
    else:
        print(f"Failed to delete page with ID {page_id}. Status code: {response.status_code}")
        print(response.text)

def retrieve_confluence_pages():
    """
    Retrieves and displays Confluence pages from a specified space using the Confluence REST API.
    Uses environment variables for authentication and configuration.
    """
    
    if not all([confluence_api_token, confluence_username, confluence_url, confluence_space_key, confluence_user_email]):
        print("Missing one or more environment variables.")
        return

    params = {
        "spaceKey": confluence_space_key,
        "expand": "body.storage,version",
        "limit": 100,
        "start": 0,
    }

    url = f"{confluence_url}/content"
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
    }

    try:
        response = requests.get(
            url,
            headers=headers,
            params=params,
            auth=HTTPBasicAuth(confluence_user_email, confluence_api_token),
        )
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Failed to retrieve pages. Error: {e}")
        return

    pages = response.json().get("results", [])
    
    page_id=[] 
    email=[]
    name=[]
    accountId=[]
    title=[]
    page_url=[]
    date=[]
    for page in pages:
        
        page_id.append(page['id'])
        email.append(page['version']['by']['email'])
        name.append(page['version']['by']['publicName'])
        accountId.append(page['version']['by']['accountId'])
        title.append(page['title'])
        date.append(page['version']['friendlyWhen'])
        page_url.append(page['_links']['webui'])
        
        
    return  page_id, email,name,accountId,title,page_url,date
def content_attachments(page_id: str) -> None:
    
    url = f"{confluence_url}/attachments"
    
    headers = {
    "Accept": "application/json"
    }
    response = requests.get(url, headers=headers, auth=AUTH)
    if response.status_code == 200:
        print("Attachments retrieved successfully!")
        attachments = response.json()["results"]
        print(f"Total attachments: {attachments}")
        for attachment in attachments:
            print(f"Attachment ID: {attachment['id']}, Title: {attachment['title']}")
    else:
        print(f"Failed to retrieve attachments. Status code: {response.status_code}")
        print(response.text)

def query_search(query):
    url = f"{confluence_url}/content/search"
    params = {
        'cql': query,
        'expand': 'body.storage'  # Include the body content in the response
    }
    try:
        response = requests.get(url, headers=HEADERS, params=params, auth=AUTH)
        response.raise_for_status()  # Raise an exception for HTTP errors
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Error making API request: {e}")
        return None

def get_page_text(response):
    page_texts = {}
    for result in response.get('results', []):
        # Extract page details directly from the result object
        page_id = result.get('id')
        page_title = result.get('title')
        body_storage = result.get('body', {}).get('storage', {}).get('value', '')
        
        if page_id and page_title:  # Ensure required fields are present
            page_texts[page_id] = {
                'title': page_title,
                'text': body_storage
            }
    return page_texts

def extract_plain_text(html):
    # Parse the HTML content
    soup = BeautifulSoup(html, 'html.parser')
    
    for macro in soup.find_all('ac:structured-macro'):
        macro.decompose()
    for tag in soup(['script', 'style', 'ac:placeholder']):
        tag.decompose()
    
    text = ""
    for element in soup.descendants:
        
        if element.find_parent(['td', 'th']):
            continue

        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            text += f"\n\n{element.get_text().strip()}\n"
        elif element.name == 'p':
            text += f"\n{element.get_text().strip()}"
        elif element.name == 'table' and not element.find_parent('table'):
            text += "\n\n[Table Start]\n"
            for row in element.find_all('tr'):
                row_text = " | ".join([cell.get_text().strip() for cell in row.find_all(['th', 'td'])])
                text += f"{row_text}\n"
            text += "[Table End]\n"
        elif element.name == 'ul' and not element.find_parent(['ul', 'ol']):
            text += "\n"
            for li in element.find_all('li'):
                text += f"- {li.get_text().strip()}\n"
        elif element.name == 'ol' and not element.find_parent(['ul', 'ol']):
            text += "\n"
            for idx, li in enumerate(element.find_all('li'), start=1):
                text += f"{idx}. {li.get_text().strip()}\n"
        elif element.name == 'pre' and not element.find_parent('pre'):
            text += f"\n\n[Code Block]\n{element.get_text().strip()}\n[End Code Block]\n"
    
    return text.strip()

def text_to_docx(plain_text, output_path):
    doc = Document()
    
    
    code_style = doc.styles.add_style('Code', 1)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.space_after = Pt(0)
    
    
    lines = plain_text.split('\n')
    
    # Track states for tables/code blocks
    in_table = False
    in_code = False
    in_tree = False
    in_sql = False
    in_json = False
    table_data = []
    code_lines = []
    tree_lines = []
    sql_lines = []
    json_lines = []
    
    for line in lines:
        line = line.strip()
        
        # Handle tables
        if line == '[Table Start]':
            in_table = True
            table_data = []
            continue
        elif line == '[Table End]':
            in_table = False
            if table_data:
                # Create table with headers
                table = doc.add_table(rows=1, cols=len(table_data[0]))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0]):
                    hdr_cells[i].text = header
                # Add table rows
                for row in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = cell
            continue
        elif in_table:
            # Split table row into cells
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                table_data.append(cells)
            continue
        
        # Handle code blocks
        if line == '[Code Block]':
            in_code = True
            code_lines = []
            continue
        elif line == '[End Code Block]':
            in_code = False
            # Add code block with style
            for code_line in code_lines:
                p = doc.add_paragraph(style='Code')
                p.add_run(code_line)
            continue
        elif in_code:
            code_lines.append(line)
            continue
        
        # Handle tree-like structures
        if any(char in line for char in ['├──', '└──']):
            if not in_tree:
                in_tree = True
                tree_lines = []
            tree_lines.append(line)
            continue
        elif in_tree:
            if not any(char in line for char in ['├──', '└──']):
                in_tree = False
                # Add tree structure as a bullet list
                for tree_line in tree_lines:
                    doc.add_paragraph(tree_line, style='List Bullet')
            continue
        
        # Handle SQL queries
        if 'SELECT' in line or 'FROM' in line or 'WHERE' in line:
            if not in_sql:
                in_sql = True
                sql_lines = []
            sql_lines.append(line)
            continue
        elif in_sql:
            if not ('SELECT' in line or 'FROM' in line or 'WHERE' in line):
                in_sql = False
                # Add SQL query as a code block
                for sql_line in sql_lines:
                    p = doc.add_paragraph(style='Code')
                    p.add_run(sql_line)
            continue
        
        # Handle JSON blocks
        if line.startswith('{') or line.startswith('['):
            if not in_json:
                in_json = True
                json_lines = []
            json_lines.append(line)
            continue
        elif in_json:
            if line.endswith('}') or line.endswith(']'):
                in_json = False
                # Add JSON block as a code block
                for json_line in json_lines:
                    p = doc.add_paragraph(style='Code')
                    p.add_run(json_line)
            continue
        
        # Handle headings (lines with no leading/trailing text)
        if not line:
            continue  # Skip empty lines
        elif all(c == '-' or c == '=' for c in line):  # Ignore Markdown-like separators
            continue
        elif line.startswith('#'):  # Markdown-style headings (optional)
            heading_level = line.count('#')
            heading_text = line.replace('#', '').strip()
            doc.add_heading(heading_text, level=min(heading_level, 3))
        else:
            # Detect lists
            if line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            else:
                # Add as normal paragraph or heading (based on context)
                is_heading = False
                if len(doc.paragraphs) > 0:
                    prev_para = doc.paragraphs[-1]
                    if prev_para.text == '' and len(line) < 100:  # Heuristic for heading
                        doc.add_heading(line, level=2)
                        is_heading = True
                if not is_heading:
                    doc.add_paragraph(line)
    
    # Save the document
    return doc
